from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer
from sqlalchemy.orm import Session
from sqlalchemy import text
from datetime import datetime, timedelta
from typing import List, Optional
import json
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
import docx
import io

from database import get_db, create_tables, User, Document, Chunk
from auth import get_password_hash, verify_password, create_access_token, get_current_user
from pydantic import BaseModel

app = FastAPI(title="Contracts SaaS API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

# Pydantic models
class UserCreate(BaseModel):
    username: str
    password: str

class UserLogin(BaseModel):
    username: str
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str

class QueryRequest(BaseModel):
    question: str

# Mock LlamaCloud response
def mock_llamacloud_parse(filename: str, content: str) -> dict:
    chunks = [
        {
            "chunk_id": "c1",
            "text": "Termination clause: Either party may terminate with 90 days' notice.",
            "embedding": [0.12, -0.45, 0.91, 0.33],
            "metadata": {"page": 2, "contract_name": filename}
        },
        {
            "chunk_id": "c2", 
            "text": "Liability cap: Limited to 12 months' fees.",
            "embedding": [0.01, 0.22, -0.87, 0.44],
            "metadata": {"page": 5, "contract_name": filename}
        },
        {
            "chunk_id": "c3",
            "text": f"Payment terms: Net 30 days from invoice date. Late fees apply after 30 days.",
            "embedding": [0.33, 0.11, 0.22, -0.55],
            "metadata": {"page": 3, "contract_name": filename}
        }
    ]
    
    return {
        "document_id": f"doc_{filename}",
        "chunks": chunks
    }

def extract_text_from_file(file: UploadFile) -> str:
    if file.content_type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.file.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif file.content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        doc = docx.Document(io.BytesIO(file.file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    else:
        return file.file.read().decode('utf-8')

@app.on_event("startup")
async def startup_event():
    try:
        create_tables()
        print("✓ Database tables created successfully")
    except Exception as e:
        print(f"⚠ Database connection failed: {e}")
        print("App will continue without database initialization")

@app.post("/signup", response_model=Token)
async def signup(user: UserCreate, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.username == user.username).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    
    hashed_password = get_password_hash(user.password)
    db_user = User(username=user.username, password_hash=hashed_password)
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    
    access_token_expires = timedelta(minutes=30)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/login", response_model=Token)
async def login(user: UserLogin, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.username == user.username).first()
    if not db_user or not verify_password(user.password, db_user.password_hash):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password"
        )
    
    access_token_expires = timedelta(minutes=30)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if file.content_type not in ["application/pdf", "text/plain", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        raise HTTPException(status_code=400, detail="Invalid file type")
    
    # Extract text content
    content = extract_text_from_file(file)
    
    # Create document record
    document = Document(
        user_id=current_user.user_id,
        filename=file.filename,
        contract_name=file.filename,
        parties="Company A, Company B",
        expiry_date=datetime.utcnow() + timedelta(days=365),
        status="Active",
        risk_score="Medium"
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    
    # Mock LlamaCloud parsing
    parsed_data = mock_llamacloud_parse(file.filename, content)
    
    # Store chunks
    for chunk_data in parsed_data["chunks"]:
        chunk = Chunk(
            doc_id=document.doc_id,
            user_id=current_user.user_id,
            text_chunk=chunk_data["text"],
            embedding=json.dumps(chunk_data["embedding"]),  # Store as JSON string
            chunk_metadata=json.dumps(chunk_data["metadata"])
        )
        db.add(chunk)
    
    db.commit()
    
    return {"message": "File uploaded and processed successfully", "doc_id": str(document.doc_id)}

@app.get("/contracts")
async def get_contracts(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    documents = db.query(Document).filter(Document.user_id == current_user.user_id).all()
    return [
        {
            "doc_id": str(doc.doc_id),
            "contract_name": doc.contract_name,
            "parties": doc.parties,
            "expiry_date": doc.expiry_date.isoformat() if doc.expiry_date else None,
            "status": doc.status,
            "risk_score": doc.risk_score,
            "uploaded_on": doc.uploaded_on.isoformat()
        }
        for doc in documents
    ]

@app.get("/contracts/{doc_id}")
async def get_contract_detail(
    doc_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    document = db.query(Document).filter(
        Document.doc_id == doc_id,
        Document.user_id == current_user.user_id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    chunks = db.query(Chunk).filter(
        Chunk.doc_id == doc_id,
        Chunk.user_id == current_user.user_id
    ).all()
    
    # Generate insights based on actual contract content
    all_text = " ".join([chunk.text_chunk for chunk in chunks])
    insights = []
    
    # Analyze actual content for insights
    if "terminate" in all_text.lower():
        insights.append({"type": "termination", "text": "Contract contains termination clauses - review notice periods"})
    if "liability" in all_text.lower():
        insights.append({"type": "liability", "text": "Liability limitations found - verify coverage adequacy"})
    if "confidential" in all_text.lower():
        insights.append({"type": "confidentiality", "text": "Confidentiality obligations present - ensure compliance"})
    if "payment" in all_text.lower():
        insights.append({"type": "payment", "text": "Payment terms specified - monitor due dates"})
    
    return {
        "doc_id": str(document.doc_id),
        "contract_name": document.contract_name,
        "parties": document.parties,
        "expiry_date": document.expiry_date.isoformat() if document.expiry_date else None,
        "status": document.status,
        "risk_score": document.risk_score,
        "uploaded_on": document.uploaded_on.isoformat(),
        "clauses": [
            {
                "title": f"Section {i+1}",
                "text": chunk.text_chunk[:150] + "..." if len(chunk.text_chunk) > 150 else chunk.text_chunk,
                "confidence": 85 + (i * 3) % 20
            }
            for i, chunk in enumerate(chunks[:5])
        ],
        "insights": insights,
        "evidence": [
            {
                "source": f"Page {json.loads(chunk.chunk_metadata).get('page', 1)}",
                "snippet": chunk.text_chunk[:200] + "..." if len(chunk.text_chunk) > 200 else chunk.text_chunk,
                "relevance": 0.9 - (i * 0.1)
            }
            for i, chunk in enumerate(chunks[:4])
        ]
    }

@app.post("/ask")
async def ask_question(
    query: QueryRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Mock query embedding
    query_embedding = np.array([0.1, 0.3, 0.5, 0.2])
    
    # Get all chunks for user
    chunks = db.query(Chunk).filter(Chunk.user_id == current_user.user_id).all()
    
    if not chunks:
        return {
            "answer": "No contracts found. Please upload some contracts first.",
            "chunks": []
        }
    
    # Calculate similarity scores
    chunk_scores = []
    for chunk in chunks:
        try:
            chunk_embedding = np.array(json.loads(chunk.embedding))
            similarity = cosine_similarity([query_embedding], [chunk_embedding])[0][0]
        except:
            similarity = 0.5  # Default similarity if parsing fails
        chunk_scores.append((chunk, similarity))
    
    # Sort by similarity and get top 3
    chunk_scores.sort(key=lambda x: x[1], reverse=True)
    top_chunks = chunk_scores[:3]
    
    return {
        "answer": f"Based on your contracts, here's what I found regarding '{query.question}': The most relevant clauses indicate specific terms and conditions that apply to your situation.",
        "chunks": [
            {
                "text": chunk.text_chunk,
                "metadata": json.loads(chunk.chunk_metadata),
                "relevance_score": round(score * 100, 1)
            }
            for chunk, score in top_chunks
        ]
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)